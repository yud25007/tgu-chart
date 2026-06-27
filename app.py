#!/usr/bin/env python3
from __future__ import annotations

import io
import json
import mimetypes
import re
import secrets
import socket
import sqlite3
import sys
import time
import webbrowser
import base64
import hashlib
import hmac
import os
from http.cookies import SimpleCookie
from copy import deepcopy
from datetime import datetime
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse
from zoneinfo import ZoneInfo

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from openpyxl import load_workbook
except ImportError as exc:
    print("缺少依赖 python-docx 或 openpyxl。请先运行：python3 -m pip install python-docx openpyxl", file=sys.stderr)
    raise SystemExit(1) from exc


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "模板.docx"
OUTPUT_DIR = Path(os.environ.get("OUTPUT_DIR", BASE_DIR / "output" / "doc"))
DATA_DIR = Path(os.environ.get("DATA_DIR", BASE_DIR / "data"))
DB_PATH = DATA_DIR / "auto_form.sqlite3"
AUTH_PASSWORD = os.environ.get("AUTH_PASSWORD", "").strip()
if not AUTH_PASSWORD:
    print("缺少 AUTH_PASSWORD 环境变量。请复制 .env.example 为 .env 并设置访问密码。", file=sys.stderr)
    raise SystemExit(1)
AUTH_SECRET = os.environ.get("AUTH_SECRET") or hashlib.sha256(f"{AUTH_PASSWORD}|{BASE_DIR}".encode("utf-8")).hexdigest()
AUTH_COOKIE_NAME = "tgu_chart_auth"
AUTH_MAX_AGE = 60 * 60 * 24 * 7
AUTH_ONLINE_WINDOW = int(os.environ.get("AUTH_ONLINE_WINDOW_SECONDS", str(15 * 60)))
APP_TIMEZONE = ZoneInfo(os.environ.get("APP_TIMEZONE", "Asia/Shanghai"))
DETAIL_START_ROW = 1
DETAIL_TEMPLATE_ROW = 3
OPINION_ROWS = 1
TITLE_TYPE_LABELS = {
    "ideology": "思政",
    "practice": "实践",
}
DEPARTMENT_SIGNATURE_TEXT = "团学部门：团委实践部"
STATIC_FILES = {
    "/": BASE_DIR / "index.html",
    "/index.html": BASE_DIR / "index.html",
    "/admin": BASE_DIR / "admin.html",
    "/admin.html": BASE_DIR / "admin.html",
    "/app.css": BASE_DIR / "app.css",
    "/app.js": BASE_DIR / "app.js",
    "/admin.js": BASE_DIR / "admin.js",
}
PUBLIC_STATIC_FILES = {"/app.css"}
LOGIN_PAGE = """<!doctype html>
<html lang="zh-CN">
  <head>
    <meta charset="utf-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <title>登录 - 思政学分记录表自动生成器</title>
    <link rel="stylesheet" href="/app.css" />
  </head>
  <body>
    <main class="login-shell">
      <section class="login-card" aria-labelledby="login-title">
        <p class="eyebrow">电气工程学院</p>
        <h1 id="login-title">表单工具登录</h1>
        <form id="login-form" autocomplete="off">
          <label for="password">访问密码</label>
          <input id="password" name="password" type="password" required autofocus />
          <button class="primary" type="submit" id="login-submit">登录</button>
        </form>
        <p class="message" id="login-message" role="status" aria-live="polite"></p>
      </section>
    </main>
    <script>
      const form = document.querySelector("#login-form");
      const submit = document.querySelector("#login-submit");
      const message = document.querySelector("#login-message");
      form.addEventListener("submit", async (event) => {
        event.preventDefault();
        message.textContent = "正在登录...";
        message.className = "message";
        submit.disabled = true;
        try {
          const password = document.querySelector("#password").value;
          const response = await fetch("/login", {
            method: "POST",
            headers: { "Content-Type": "application/json" },
            body: JSON.stringify({ password }),
          });
          const result = await response.json();
          if (!response.ok || !result.ok) throw new Error(result.message || "登录失败");
          const next = new URLSearchParams(window.location.search).get("next") || "/";
          window.location.href = next.startsWith("/") && !next.startsWith("//") ? next : "/";
        } catch (error) {
          message.textContent = error.message;
          message.className = "message error";
        } finally {
          submit.disabled = false;
        }
      });
    </script>
  </body>
</html>
"""


class ValidationError(ValueError):
    pass


def now_text() -> str:
    return datetime.now(APP_TIMEZONE).strftime("%Y-%m-%d %H:%M:%S")


def db_connect() -> sqlite3.Connection:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_db() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    with db_connect() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                activity_name TEXT NOT NULL,
                activity_date TEXT NOT NULL,
                title_type TEXT NOT NULL DEFAULT 'ideology',
                credit_type TEXT NOT NULL,
                default_points TEXT NOT NULL DEFAULT '',
                entry_count INTEGER NOT NULL,
                output_filename TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS record_entries (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                record_id INTEGER NOT NULL REFERENCES records(id) ON DELETE CASCADE,
                sort_order INTEGER NOT NULL,
                class_name TEXT NOT NULL,
                name TEXT NOT NULL DEFAULT '',
                student_id TEXT NOT NULL,
                points TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS login_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ip TEXT NOT NULL,
                user_agent TEXT NOT NULL DEFAULT '',
                success INTEGER NOT NULL,
                session_id TEXT NOT NULL DEFAULT '',
                created_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS auth_sessions (
                session_id TEXT PRIMARY KEY,
                ip TEXT NOT NULL,
                user_agent TEXT NOT NULL DEFAULT '',
                created_at TEXT NOT NULL,
                last_seen_at TEXT NOT NULL,
                expires_at INTEGER NOT NULL,
                logged_out_at TEXT
            );

            CREATE INDEX IF NOT EXISTS idx_records_created_at ON records(created_at DESC);
            CREATE INDEX IF NOT EXISTS idx_record_entries_record_id ON record_entries(record_id, sort_order);
            CREATE INDEX IF NOT EXISTS idx_login_events_created_at ON login_events(created_at DESC);
            CREATE INDEX IF NOT EXISTS idx_auth_sessions_online ON auth_sessions(logged_out_at, expires_at, last_seen_at);
            """
        )
        ensure_record_schema(conn)


def ensure_record_schema(conn: sqlite3.Connection) -> None:
    columns = {row["name"] for row in conn.execute("PRAGMA table_info(records)").fetchall()}
    if "title_type" not in columns:
        conn.execute("ALTER TABLE records ADD COLUMN title_type TEXT NOT NULL DEFAULT 'ideology'")
        conn.execute("UPDATE records SET title_type = 'practice' WHERE credit_type LIKE '%实践%'")


def first_free_port(start: int = 8765, host: str = "127.0.0.1") -> int:
    for port in range(start, start + 50):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
            sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            try:
                sock.bind((host, port))
            except OSError:
                continue
            return port
    raise RuntimeError("没有找到可用端口")


def require_text(payload: dict, key: str, label: str, max_length: int = 80) -> str:
    value = str(payload.get(key, "")).strip()
    if not value:
        raise ValidationError(f"请填写{label}")
    if len(value) > max_length:
        raise ValidationError(f"{label}不能超过 {max_length} 个字符")
    return value


def clean_optional_text(value: object, max_length: int = 80) -> str:
    text = str(value or "").strip()
    if len(text) > max_length:
        raise ValidationError(f"字段内容不能超过 {max_length} 个字符")
    return text


def infer_title_type(credit_type: str) -> str:
    return "practice" if "实践" in credit_type else "ideology"


def normalize_title_type(value: object, credit_type: str) -> str:
    text = clean_optional_text(value, 20)
    aliases = {
        "ideology": "ideology",
        "思政": "ideology",
        "思想政治": "ideology",
        "思政学分": "ideology",
        "practice": "practice",
        "实践": "practice",
        "实践学分": "practice",
    }
    if not text:
        return infer_title_type(credit_type)
    return aliases.get(text, infer_title_type(credit_type))


def document_title_text(title_type: str) -> str:
    return f"电气工程学院{TITLE_TYPE_LABELS.get(title_type, TITLE_TYPE_LABELS['ideology'])}学分记录表"


def normalize_entry(entry: dict, index: int, default_points: str) -> dict:
    class_name = clean_optional_text(entry.get("className"), 40)
    name = clean_optional_text(entry.get("name"), 40)
    student_id = clean_optional_text(entry.get("studentId"), 40)
    points = clean_optional_text(entry.get("points"), 20) or default_points

    if not class_name and not name and not student_id and not points:
        raise ValidationError(f"第 {index} 行是空行")
    if not class_name:
        raise ValidationError(f"请填写第 {index} 行班级")
    if not student_id:
        raise ValidationError(f"请填写第 {index} 行学号")
    if not points:
        raise ValidationError(f"请填写第 {index} 行加分数量")

    return {
        "class_name": class_name,
        "name": name,
        "student_id": student_id,
        "points": points,
    }


def ensure_no_duplicate_student_ids(entries: list[dict]) -> None:
    seen: dict[str, int] = {}
    for index, entry in enumerate(entries, start=1):
        student_id = entry["student_id"]
        if student_id in seen:
            raise ValidationError(f"第 {seen[student_id]} 行和第 {index} 行学号重复：{student_id}")
        seen[student_id] = index


def normalize_payload(payload: dict) -> dict:
    default_points = clean_optional_text(payload.get("defaultPoints"), 20)
    entries_payload = payload.get("entries")

    if isinstance(entries_payload, list):
        non_empty_entries = [
            entry
            for entry in entries_payload
            if isinstance(entry, dict)
            and any(str(entry.get(key, "")).strip() for key in ("className", "studentId", "points"))
        ]
    else:
        non_empty_entries = []

    if not non_empty_entries:
        non_empty_entries = [
            {
                "className": payload.get("className"),
                "name": payload.get("name"),
                "studentId": payload.get("studentId"),
                "points": payload.get("points") or default_points,
            }
        ]

    if len(non_empty_entries) > 120:
        raise ValidationError("一次最多生成 120 条明细")

    entries = [
        normalize_entry(entry, index + 1, default_points)
        for index, entry in enumerate(non_empty_entries)
    ]
    ensure_no_duplicate_student_ids(entries)

    credit_type = require_text(payload, "creditType", "加分类型", 40)

    return {
        "activity_name": require_text(payload, "activityName", "活动名称", 120),
        "date": require_text(payload, "date", "日期", 40),
        "default_points": default_points,
        "entries": entries,
        "credit_type": credit_type,
        "title_type": normalize_title_type(payload.get("titleType"), credit_type),
    }


def display_date(raw: str) -> str:
    try:
        parsed = datetime.strptime(raw, "%Y-%m-%d")
    except ValueError:
        return raw
    return f"{parsed.year}年{parsed.month}月{parsed.day}日"


def safe_filename_part(text: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\\s]+', "_", text.strip())
    cleaned = cleaned.strip("._")
    return cleaned[:40] or "未命名"


def remove_run(run) -> None:
    run._element.getparent().remove(run._element)


def trim_to_one_paragraph(cell):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for extra in list(cell.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)
    return paragraph


def paragraph_run_properties(paragraph):
    p_pr = paragraph._p.pPr
    return p_pr.find(qn("w:rPr")) if p_pr is not None else None


def run_properties(run):
    return run._element.find(qn("w:rPr"))


def apply_run_properties(run, source_r_pr) -> None:
    if source_r_pr is None:
        return

    existing = run._element.find(qn("w:rPr"))
    if existing is not None:
        run._element.remove(existing)
    run._element.insert(0, deepcopy(source_r_pr))


def set_cell_text(cell, text: str, *, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    paragraph = trim_to_one_paragraph(cell)
    paragraph.alignment = align
    runs = list(paragraph.runs)

    if runs:
        source_r_pr = run_properties(runs[0])
        run = runs[0]
        run.text = text
        for extra in runs[1:]:
            remove_run(extra)
    elif text:
        source_r_pr = paragraph_run_properties(paragraph)
        run = paragraph.add_run(text)
        apply_run_properties(run, source_r_pr)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_paragraph_run_texts(paragraph, texts: list[str]) -> None:
    runs = list(paragraph.runs)
    fallback_r_pr = paragraph_run_properties(paragraph)

    for index, text in enumerate(texts):
        if index < len(runs):
            run = runs[index]
        else:
            run = paragraph.add_run()
            apply_run_properties(run, fallback_r_pr)
        run.text = text

    for extra in runs[len(texts):]:
        remove_run(extra)


def set_activity_header(cell, activity_name: str, date_text: str) -> None:
    paragraph = trim_to_one_paragraph(cell)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_run_texts(
        paragraph,
        [
            "活动名称",
            ":",
            f" {activity_name}    ",
            "时间：",
            f" {date_text}",
        ],
    )
    runs = paragraph.runs
    if len(runs) >= 5:
        source_r_pr = run_properties(runs[2])
        if source_r_pr is None:
            source_r_pr = run_properties(runs[3])
        apply_run_properties(runs[4], source_r_pr)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_document_title(doc, title_type: str) -> None:
    title_paragraph = next(
        (
            paragraph
            for paragraph in doc.paragraphs
            if "电气工程学院" in paragraph.text and "记录表" in paragraph.text
        ),
        None,
    )
    if title_paragraph is None:
        return

    set_paragraph_run_texts(
        title_paragraph,
        [
            "电气工程",
            "学院",
            TITLE_TYPE_LABELS.get(title_type, TITLE_TYPE_LABELS["ideology"]),
            "学分",
            "记录表",
        ],
    )


def remove_table_row(table, row_index: int) -> None:
    row_element = table.rows[row_index]._tr
    row_element.getparent().remove(row_element)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for existing in tr_pr.findall(qn("w:tblHeader")):
        tr_pr.remove(existing)
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def ensure_detail_rows(table, entry_count: int) -> None:
    current_capacity = len(table.rows) - DETAIL_START_ROW - OPINION_ROWS
    if current_capacity < 1:
        raise RuntimeError("模板表格缺少明细行")

    while current_capacity < entry_count:
        new_row = deepcopy(table.rows[DETAIL_TEMPLATE_ROW]._tr)
        table.rows[-1]._tr.addprevious(new_row)
        current_capacity += 1


def clear_detail_rows(table) -> None:
    last_detail_row = len(table.rows) - OPINION_ROWS
    for row_index in range(DETAIL_START_ROW, last_detail_row):
        for col_index in (1, 3, 5, 7):
            set_cell_text(table.cell(row_index, col_index), "")


def credit_value_text(points: str, credit_type: str) -> str:
    return f"{points}{credit_type}"


def fill_detail_row(table, row_index: int, entry: dict, credit_type: str) -> None:
    set_cell_text(table.cell(row_index, 1), entry.get("name", ""))
    set_cell_text(table.cell(row_index, 3), entry["class_name"])
    set_cell_text(table.cell(row_index, 5), entry["student_id"])
    set_cell_text(table.cell(row_index, 7), credit_value_text(entry["points"], credit_type))


def paragraph_has_visible_text(paragraph) -> bool:
    return bool(paragraph.text.strip())


def empty_paragraph_like(paragraph):
    new_paragraph = deepcopy(paragraph._p)
    for run in list(new_paragraph.findall(qn("w:r"))):
        new_paragraph.remove(run)
    return new_paragraph


def ensure_leading_blank_paragraphs(cell, count: int) -> None:
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        return

    first_content_index = next(
        (index for index, paragraph in enumerate(paragraphs) if paragraph_has_visible_text(paragraph)),
        len(paragraphs),
    )
    missing_count = count - first_content_index
    if missing_count <= 0:
        return

    template = paragraphs[0]
    insert_before = paragraphs[first_content_index]._p if first_content_index < len(paragraphs) else None
    for _ in range(missing_count):
        new_paragraph = empty_paragraph_like(template)
        if insert_before is not None:
            insert_before.addprevious(new_paragraph)
        else:
            cell._tc.append(new_paragraph)


def paragraph_contains_text(paragraph, text: str) -> bool:
    return text in paragraph.text


def clone_paragraph_with_text(paragraph, text: str):
    new_paragraph = deepcopy(paragraph._p)
    text_elements = list(new_paragraph.iter(qn("w:t")))
    if text_elements:
        text_elements[0].text = text
        for extra in text_elements[1:]:
            extra.text = ""
    return new_paragraph


def ensure_department_signature_line(cell) -> None:
    paragraphs = list(cell.paragraphs)
    if any(paragraph_contains_text(paragraph, "团学部门") for paragraph in paragraphs):
        return

    signature_paragraph = next(
        (paragraph for paragraph in paragraphs if paragraph_contains_text(paragraph, "辅导员签字")),
        None,
    )
    if signature_paragraph is None:
        return

    department_paragraph = clone_paragraph_with_text(signature_paragraph, DEPARTMENT_SIGNATURE_TEXT)
    signature_paragraph._p.addprevious(department_paragraph)


def ensure_opinion_signature_spacing(table) -> None:
    opinion_row = table.rows[-1]
    signature_cell = opinion_row.cells[1]
    ensure_leading_blank_paragraphs(signature_cell, 2)
    ensure_department_signature_line(signature_cell)


def element_has_visible_content(element) -> bool:
    for text_element in element.iter(qn("w:t")):
        if text_element.text and text_element.text.strip():
            return True
    return any(
        child.tag in {qn("w:drawing"), qn("w:pict"), qn("w:object")}
        for child in element.iter()
    )


def remove_empty_paragraphs_after_table(doc, table) -> None:
    body = doc._body._element
    children = list(body)
    try:
        table_index = children.index(table._tbl)
    except ValueError:
        return

    for element in children[table_index + 1:]:
        if element.tag != qn("w:p"):
            continue
        if element_has_visible_content(element):
            continue
        body.remove(element)


def cell_to_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def normalize_header(text: str) -> str:
    return re.sub(r"\s+", "", text).lower()


def find_header_map(rows: list[list[str]]) -> tuple[int, dict[str, int]]:
    aliases = {
        "className": ("班级", "班别", "班级名称"),
        "name": ("姓名", "名字", "学生姓名"),
        "studentId": ("学号", "学生号", "学籍号"),
        "points": ("加分", "加分数量", "学分", "所获学分", "分数"),
    }

    for row_index, row in enumerate(rows[:12]):
        header_map: dict[str, int] = {}
        normalized = [normalize_header(value) for value in row]
        for field, names in aliases.items():
            for col_index, value in enumerate(normalized):
                if any(normalize_header(name) in value for name in names):
                    header_map[field] = col_index
                    break
        if "className" in header_map and "studentId" in header_map:
            return row_index, header_map

    return -1, {}


def looks_like_student_id(value: str) -> bool:
    return bool(re.fullmatch(r"[A-Za-z]?\d{6,}[A-Za-z]?", value.strip()))


def infer_entry_from_row(row: list[str]) -> dict:
    entry = {"className": "", "name": "", "studentId": "", "points": ""}
    for value in row:
        if not value:
            continue
        if not entry["studentId"] and looks_like_student_id(value):
            entry["studentId"] = value
        elif not entry["className"] and ("班" in value or "电气" in value) and any(ch.isdigit() for ch in value):
            entry["className"] = value
        elif not entry["points"] and re.fullmatch(r"\d+(?:\.\d+)?", value):
            entry["points"] = value
        elif not entry["name"]:
            entry["name"] = value
    return entry


def infer_metadata_from_filename(file_name: str) -> dict:
    stem = Path(file_name or "").stem
    inferred = {"activityName": "", "date": ""}
    if not stem:
        return inferred

    date_match = re.search(r"(?:(20\d{2})[.\-_/年])?(\d{1,2})[.\-_/月](\d{1,2})日?", stem)
    activity_name = stem
    if date_match:
        year = int(date_match.group(1) or datetime.now(APP_TIMEZONE).year)
        month = int(date_match.group(2))
        day = int(date_match.group(3))
        try:
            inferred["date"] = datetime(year, month, day).strftime("%Y-%m-%d")
        except ValueError:
            inferred["date"] = ""
        activity_name = stem[: date_match.start()] + stem[date_match.end() :]

    activity_name = re.sub(r"(活动)?统计表$|名单$|活动名单$|信息表$", "", activity_name)
    activity_name = re.sub(r"^[\s._\-—-]+|[\s._\-—-]+$", "", activity_name)
    inferred["activityName"] = activity_name
    return inferred


def parse_workbook_entries(file_bytes: bytes, file_name: str = "") -> dict:
    workbook = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    worksheet = workbook.active
    rows = [
        [cell_to_text(cell) for cell in row]
        for row in worksheet.iter_rows(values_only=True)
    ]
    rows = [row for row in rows if any(row)]

    if not rows:
        raise ValidationError("Excel 里没有可识别的数据")

    header_index, header_map = find_header_map(rows)
    data_rows = rows[header_index + 1:] if header_index >= 0 else rows
    entries = []
    warnings = []

    for source_row_number, row in enumerate(data_rows, start=(header_index + 2 if header_index >= 0 else 1)):
        if header_map:
            def mapped_value(field: str) -> str:
                col_index = header_map.get(field)
                if col_index is None or col_index >= len(row):
                    return ""
                return row[col_index]

            entry = {
                "className": mapped_value("className"),
                "name": mapped_value("name"),
                "studentId": mapped_value("studentId"),
                "points": mapped_value("points"),
            }
        else:
            entry = infer_entry_from_row(row)

        if not any(entry.values()):
            continue

        entries.append(entry)
        missing = [label for label, key in (("班级", "className"), ("学号", "studentId")) if not entry[key]]
        if missing:
            warnings.append(f"Excel 第 {source_row_number} 行缺少{'、'.join(missing)}")

    if not entries:
        raise ValidationError("没有识别到名单明细")

    seen_student_ids: dict[str, int] = {}
    for index, entry in enumerate(entries, start=1):
        student_id = entry.get("studentId", "")
        if not student_id:
            continue
        if student_id in seen_student_ids:
            warnings.append(f"第 {seen_student_ids[student_id]} 行和第 {index} 行学号重复：{student_id}")
        else:
            seen_student_ids[student_id] = index

    return {
        "sheetName": worksheet.title,
        "entries": entries,
        "warnings": warnings,
        "headerDetected": bool(header_map),
        "columns": header_map,
        "inferred": infer_metadata_from_filename(file_name),
    }


def build_docx(data: dict) -> tuple[bytes, str, Path]:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"找不到模板文件：{TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    if not doc.tables:
        raise RuntimeError("模板中没有表格，无法自动填写")

    table = doc.tables[0]
    if len(table.rows) < 6 or len(table.columns) < 8:
        raise RuntimeError("模板表格结构与预期不一致，至少需要 6 行 8 列")

    set_document_title(doc, data["title_type"])
    set_activity_header(table.cell(0, 0), data["activity_name"], display_date(data["date"]))
    remove_table_row(table, 1)
    set_repeat_table_header(table.rows[0])

    ensure_detail_rows(table, len(data["entries"]))
    clear_detail_rows(table)

    for offset, entry in enumerate(data["entries"]):
        fill_detail_row(table, DETAIL_START_ROW + offset, entry, data["credit_type"])

    ensure_opinion_signature_spacing(table)
    remove_empty_paragraphs_after_table(doc, table)

    first_entry = data["entries"][0]

    filename = "_".join(
        [
            safe_filename_part(data["date"]),
            safe_filename_part(data["activity_name"]),
            safe_filename_part(first_entry["class_name"]),
            f"{len(data['entries'])}人",
        ]
    )
    filename = f"{filename}.docx"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / filename
    doc.save(output_path)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), filename, output_path


def record_to_client(row: sqlite3.Row) -> dict:
    title_type = row["title_type"] if "title_type" in row.keys() else infer_title_type(row["credit_type"])
    return {
        "id": row["id"],
        "activityName": row["activity_name"],
        "date": row["activity_date"],
        "titleType": title_type,
        "recordTitle": document_title_text(title_type),
        "creditType": row["credit_type"],
        "defaultPoints": row["default_points"],
        "entryCount": row["entry_count"],
        "outputFilename": row["output_filename"],
        "createdAt": row["created_at"],
        "updatedAt": row["updated_at"],
    }


def entry_to_client(row: sqlite3.Row) -> dict:
    return {
        "className": row["class_name"],
        "name": row["name"],
        "studentId": row["student_id"],
        "points": row["points"],
    }


def save_generation_record(data: dict, filename: str) -> int:
    stamp = now_text()
    with db_connect() as conn:
        cursor = conn.execute(
            """
            INSERT INTO records (
                activity_name, activity_date, title_type, credit_type, default_points,
                entry_count, output_filename, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                data["activity_name"],
                data["date"],
                data["title_type"],
                data["credit_type"],
                data.get("default_points", ""),
                len(data["entries"]),
                filename,
                stamp,
                stamp,
            ),
        )
        record_id = int(cursor.lastrowid)
        conn.executemany(
            """
            INSERT INTO record_entries (
                record_id, sort_order, class_name, name, student_id, points
            )
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            [
                (
                    record_id,
                    index,
                    entry["class_name"],
                    entry.get("name", ""),
                    entry["student_id"],
                    entry["points"],
                )
                for index, entry in enumerate(data["entries"])
            ],
        )
    return record_id


def list_records(limit: int = 30) -> list[dict]:
    limit = max(1, min(limit, 100))
    with db_connect() as conn:
        rows = conn.execute(
            """
            SELECT * FROM records
            ORDER BY created_at DESC, id DESC
            LIMIT ?
            """,
            (limit,),
        ).fetchall()
    return [record_to_client(row) for row in rows]


def get_record(record_id: int) -> dict | None:
    with db_connect() as conn:
        record = conn.execute("SELECT * FROM records WHERE id = ?", (record_id,)).fetchone()
        if record is None:
            return None
        entries = conn.execute(
            """
            SELECT * FROM record_entries
            WHERE record_id = ?
            ORDER BY sort_order ASC, id ASC
            """,
            (record_id,),
        ).fetchall()

    payload = record_to_client(record)
    payload["entries"] = [entry_to_client(row) for row in entries]
    return payload


def delete_record(record_id: int) -> bool:
    record = get_record(record_id)
    if record is None:
        return False

    with db_connect() as conn:
        conn.execute("DELETE FROM records WHERE id = ?", (record_id,))

    output_path = OUTPUT_DIR / record["outputFilename"]
    if output_path.exists():
        try:
            output_path.unlink()
        except OSError:
            pass
    return True


def delete_records(record_ids: list[int]) -> dict:
    unique_ids = list(dict.fromkeys(record_ids))
    if not unique_ids:
        return {"deletedIds": [], "missingIds": [], "deletedCount": 0}

    placeholders = ",".join("?" for _ in unique_ids)
    with db_connect() as conn:
        rows = conn.execute(
            f"""
            SELECT id, output_filename
            FROM records
            WHERE id IN ({placeholders})
            """,
            unique_ids,
        ).fetchall()
        found_ids = {int(row["id"]) for row in rows}
        conn.execute(
            f"DELETE FROM records WHERE id IN ({placeholders})",
            unique_ids,
        )

    for row in rows:
        output_path = OUTPUT_DIR / row["output_filename"]
        if output_path.exists():
            try:
                output_path.unlink()
            except OSError:
                pass

    deleted_ids = [record_id for record_id in unique_ids if record_id in found_ids]
    missing_ids = [record_id for record_id in unique_ids if record_id not in found_ids]
    return {
        "deletedIds": deleted_ids,
        "missingIds": missing_ids,
        "deletedCount": len(deleted_ids),
    }


def normalize_record_ids(value: object, limit: int = 100) -> list[int]:
    if not isinstance(value, list):
        raise ValidationError("请选择要删除的历史记录")
    if len(value) > limit:
        raise ValidationError(f"一次最多删除 {limit} 条历史记录")

    record_ids: list[int] = []
    for index, item in enumerate(value, start=1):
        try:
            record_id = int(item)
        except (TypeError, ValueError):
            raise ValidationError(f"第 {index} 个记录编号无效") from None
        if record_id <= 0:
            raise ValidationError(f"第 {index} 个记录编号无效")
        record_ids.append(record_id)

    if not record_ids:
        raise ValidationError("请选择要删除的历史记录")
    return record_ids


def unix_to_text(timestamp: int | float) -> str:
    return datetime.fromtimestamp(timestamp, APP_TIMEZONE).strftime("%Y-%m-%d %H:%M:%S")


def compact_header_value(value: str, max_length: int = 240) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text[:max_length]


def sign_auth_value(session_id: str, expires_at: int) -> str:
    message = f"{session_id}:{expires_at}".encode("utf-8")
    return hmac.new(AUTH_SECRET.encode("utf-8"), message, hashlib.sha256).hexdigest()


def create_auth_cookie_value(session_id: str, expires_at: int) -> str:
    return f"{session_id}:{expires_at}:{sign_auth_value(session_id, expires_at)}"


def parse_auth_cookie_value(value: str) -> tuple[str, int] | None:
    try:
        session_id, expires_text, signature = value.split(":", 2)
        expires_at = int(expires_text)
    except (TypeError, ValueError):
        return None
    if not session_id or len(session_id) > 120:
        return None
    if expires_at < int(time.time()):
        return None
    if not hmac.compare_digest(signature, sign_auth_value(session_id, expires_at)):
        return None
    return session_id, expires_at


def create_auth_session(ip: str, user_agent: str) -> tuple[str, int]:
    session_id = secrets.token_urlsafe(32)
    expires_at = int(time.time()) + AUTH_MAX_AGE
    stamp = now_text()
    with db_connect() as conn:
        conn.execute(
            """
            INSERT INTO auth_sessions (
                session_id, ip, user_agent, created_at, last_seen_at, expires_at
            )
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (session_id, ip, user_agent, stamp, stamp, expires_at),
        )
    return session_id, expires_at


def record_login_event(ip: str, user_agent: str, success: bool, session_id: str = "") -> None:
    with db_connect() as conn:
        conn.execute(
            """
            INSERT INTO login_events (ip, user_agent, success, session_id, created_at)
            VALUES (?, ?, ?, ?, ?)
            """,
            (ip, user_agent, 1 if success else 0, session_id, now_text()),
        )


def mark_session_logged_out(session_id: str) -> None:
    with db_connect() as conn:
        conn.execute(
            """
            UPDATE auth_sessions
            SET logged_out_at = COALESCE(logged_out_at, ?), last_seen_at = ?
            WHERE session_id = ?
            """,
            (now_text(), now_text(), session_id),
        )


def active_session_from_cookie_value(value: str) -> str | None:
    parsed = parse_auth_cookie_value(value)
    if parsed is None:
        return None
    session_id, expires_at = parsed
    with db_connect() as conn:
        row = conn.execute(
            """
            SELECT session_id, expires_at
            FROM auth_sessions
            WHERE session_id = ?
              AND logged_out_at IS NULL
              AND expires_at >= ?
            """,
            (session_id, int(time.time())),
        ).fetchone()
        if row is None or int(row["expires_at"]) != expires_at:
            return None
        conn.execute(
            "UPDATE auth_sessions SET last_seen_at = ? WHERE session_id = ?",
            (now_text(), session_id),
        )
    return session_id


def client_ip_from_headers(headers, fallback: str) -> str:
    candidates = [
        headers.get("CF-Connecting-IP"),
        headers.get("X-Real-IP"),
    ]
    forwarded_for = headers.get("X-Forwarded-For", "")
    if forwarded_for:
        candidates.append(forwarded_for.split(",", 1)[0])
    candidates.append(fallback)

    for candidate in candidates:
        ip = compact_header_value(candidate, 80)
        if ip:
            return ip
    return "unknown"


def online_cutoff_text() -> str:
    return unix_to_text(int(time.time()) - AUTH_ONLINE_WINDOW)


def admin_dashboard_data() -> dict:
    now_ts = int(time.time())
    cutoff = online_cutoff_text()
    day_cutoff = unix_to_text(now_ts - 24 * 60 * 60)

    with db_connect() as conn:
        online_rows = conn.execute(
            """
            SELECT session_id, ip, user_agent, created_at, last_seen_at, expires_at
            FROM auth_sessions
            WHERE logged_out_at IS NULL
              AND expires_at >= ?
              AND last_seen_at >= ?
            ORDER BY last_seen_at DESC, created_at DESC
            """,
            (now_ts, cutoff),
        ).fetchall()
        event_rows = conn.execute(
            """
            SELECT ip, user_agent, success, created_at
            FROM login_events
            ORDER BY created_at DESC, id DESC
            LIMIT 60
            """
        ).fetchall()
        counters = conn.execute(
            """
            SELECT
                SUM(CASE WHEN success = 1 AND created_at >= ? THEN 1 ELSE 0 END) AS success_24h,
                SUM(CASE WHEN success = 0 AND created_at >= ? THEN 1 ELSE 0 END) AS failed_24h
            FROM login_events
            """,
            (day_cutoff, day_cutoff),
        ).fetchone()

    unique_ips = sorted({row["ip"] for row in online_rows})
    return {
        "ok": True,
        "generatedAt": now_text(),
        "onlineWindowSeconds": AUTH_ONLINE_WINDOW,
        "onlineCount": len(online_rows),
        "onlineUniqueIpCount": len(unique_ips),
        "onlineUniqueIps": unique_ips,
        "success24h": int(counters["success_24h"] or 0),
        "failed24h": int(counters["failed_24h"] or 0),
        "onlineSessions": [
            {
                "ip": row["ip"],
                "userAgent": row["user_agent"],
                "loginAt": row["created_at"],
                "lastSeenAt": row["last_seen_at"],
                "expiresAt": unix_to_text(row["expires_at"]),
            }
            for row in online_rows
        ],
        "recentLogins": [
            {
                "ip": row["ip"],
                "userAgent": row["user_agent"],
                "success": bool(row["success"]),
                "createdAt": row["created_at"],
            }
            for row in event_rows
        ],
    }


class AutoFormHandler(SimpleHTTPRequestHandler):
    server_version = "AutoFormServer/1.0"

    def log_message(self, fmt: str, *args) -> None:
        print(f"[网页表单] {self.address_string()} - {fmt % args}")

    def is_secure_request(self) -> bool:
        return self.headers.get("X-Forwarded-Proto", "").lower() == "https"

    def client_ip(self) -> str:
        fallback = self.client_address[0] if self.client_address else ""
        return client_ip_from_headers(self.headers, fallback)

    def user_agent(self) -> str:
        return compact_header_value(self.headers.get("User-Agent", ""))

    def auth_cookie_value(self) -> str:
        cookie_header = self.headers.get("Cookie", "")
        if not cookie_header:
            return ""
        cookies = SimpleCookie()
        try:
            cookies.load(cookie_header)
        except Exception:
            return ""
        morsel = cookies.get(AUTH_COOKIE_NAME)
        return morsel.value if morsel is not None else ""

    def is_authenticated(self) -> bool:
        if hasattr(self, "_authenticated"):
            return bool(self._authenticated)
        session_id = active_session_from_cookie_value(self.auth_cookie_value())
        self._auth_session_id = session_id or ""
        self._authenticated = bool(session_id)
        return bool(session_id)

    def auth_cookie_header(self, session_id: str, expires_at: int) -> str:
        parts = [
            f"{AUTH_COOKIE_NAME}={create_auth_cookie_value(session_id, expires_at)}",
            "Path=/",
            f"Max-Age={AUTH_MAX_AGE}",
            "HttpOnly",
            "SameSite=Lax",
        ]
        if self.is_secure_request():
            parts.append("Secure")
        return "; ".join(parts)

    def expire_current_session(self) -> None:
        parsed = parse_auth_cookie_value(self.auth_cookie_value())
        if parsed is None:
            return
        mark_session_logged_out(parsed[0])

    def clear_auth_cookie_header(self) -> str:
        parts = [
            f"{AUTH_COOKIE_NAME}=",
            "Path=/",
            "Max-Age=0",
            "HttpOnly",
            "SameSite=Lax",
        ]
        if self.is_secure_request():
            parts.append("Secure")
        return "; ".join(parts)

    def send_redirect(self, location: str) -> None:
        self.send_response(HTTPStatus.SEE_OTHER)
        self.send_header("Location", location)
        self.send_header("Content-Length", "0")
        self.end_headers()

    def send_login_page(self) -> None:
        body = LOGIN_PAGE.encode("utf-8")
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def send_auth_required(self, *, as_json: bool = True) -> None:
        if as_json:
            self.send_json({"ok": False, "message": "请先登录"}, status=HTTPStatus.UNAUTHORIZED)
        else:
            parsed = urlparse(self.path)
            next_path = parsed.path or "/"
            if parsed.query:
                next_path = f"{next_path}?{parsed.query}"
            self.send_redirect(f"/login?next={quote(next_path)}")

    def send_static_file(self, file_path: Path) -> None:
        content_type, _ = mimetypes.guess_type(file_path.name)
        body = file_path.read_bytes()
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", content_type or "application/octet-stream")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = unquote(parsed.path)
        if path == "/health":
            self.send_json({"ok": True})
            return
        if path == "/login":
            if self.is_authenticated():
                self.send_redirect("/")
            else:
                self.send_login_page()
            return
        if path in {"/favicon.ico", "/apple-touch-icon.png", "/apple-touch-icon-precomposed.png"}:
            self.send_response(HTTPStatus.NO_CONTENT)
            self.end_headers()
            return
        if path in PUBLIC_STATIC_FILES:
            file_path = STATIC_FILES.get(path)
            if file_path and file_path.exists():
                self.send_static_file(file_path)
                return

        if not self.is_authenticated():
            self.send_auth_required(as_json=path.startswith(("/history", "/record/", "/admin-data")))
            return

        if path == "/admin-data":
            self.send_json(admin_dashboard_data())
            return
        if path == "/history":
            self.send_json({"ok": True, "records": list_records()})
            return
        if path.startswith("/record/"):
            record_id = self.record_id_from_path(path)
            if record_id is None:
                self.send_json({"ok": False, "message": "记录编号无效"}, status=HTTPStatus.BAD_REQUEST)
                return
            record = get_record(record_id)
            if record is None:
                self.send_json({"ok": False, "message": "记录不存在"}, status=HTTPStatus.NOT_FOUND)
                return
            self.send_json({"ok": True, "record": record})
            return
        if path.startswith("/download/"):
            record_id = self.record_id_from_path(path)
            if record_id is None:
                self.send_error(HTTPStatus.BAD_REQUEST, "Bad record id")
                return
            record = get_record(record_id)
            if record is None:
                self.send_error(HTTPStatus.NOT_FOUND, "Record not found")
                return
            self.send_docx_file(OUTPUT_DIR / record["outputFilename"], record["outputFilename"])
            return
        file_path = STATIC_FILES.get(path)
        if not file_path or not file_path.exists():
            self.send_error(HTTPStatus.NOT_FOUND, "File not found")
            return

        self.send_static_file(file_path)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/login":
            self.handle_login()
            return
        if parsed.path == "/logout":
            self.handle_logout()
            return
        if not self.is_authenticated():
            self.send_auth_required()
            return
        if parsed.path == "/import-roster":
            self.handle_import_roster()
            return
        if parsed.path == "/records/batch-delete":
            self.handle_batch_delete_records()
            return

        if parsed.path != "/generate":
            self.send_error(HTTPStatus.NOT_FOUND, "File not found")
            return

        try:
            body_size = int(self.headers.get("Content-Length", "0"))
            if body_size > 1024 * 1024:
                raise ValidationError("提交内容过大")
            payload = json.loads(self.rfile.read(body_size).decode("utf-8"))
            data = normalize_payload(payload)
            docx_bytes, filename, _output_path = build_docx(data)
            record_id = save_generation_record(data, filename)
        except ValidationError as exc:
            self.send_json({"ok": False, "message": str(exc)}, status=HTTPStatus.BAD_REQUEST)
            return
        except Exception as exc:
            self.send_json({"ok": False, "message": f"生成失败：{exc}"}, status=HTTPStatus.INTERNAL_SERVER_ERROR)
            return

        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(filename)}")
        self.send_header("X-Record-Id", str(record_id))
        self.send_header("Content-Length", str(len(docx_bytes)))
        self.end_headers()
        self.wfile.write(docx_bytes)

    def do_DELETE(self) -> None:
        parsed = urlparse(self.path)
        path = unquote(parsed.path)
        if not self.is_authenticated():
            self.send_auth_required()
            return
        if not path.startswith("/record/"):
            self.send_error(HTTPStatus.NOT_FOUND, "File not found")
            return
        record_id = self.record_id_from_path(path)
        if record_id is None:
            self.send_json({"ok": False, "message": "记录编号无效"}, status=HTTPStatus.BAD_REQUEST)
            return
        if not delete_record(record_id):
            self.send_json({"ok": False, "message": "记录不存在"}, status=HTTPStatus.NOT_FOUND)
            return
        self.send_json({"ok": True})

    def handle_batch_delete_records(self) -> None:
        try:
            body = self.read_request_body(64 * 1024)
            payload = json.loads(body.decode("utf-8") or "{}")
            record_ids = normalize_record_ids(payload.get("recordIds"))
            result = delete_records(record_ids)
        except ValidationError as exc:
            self.send_json({"ok": False, "message": str(exc)}, status=HTTPStatus.BAD_REQUEST)
            return
        except Exception as exc:
            self.send_json({"ok": False, "message": f"批量删除失败：{exc}"}, status=HTTPStatus.INTERNAL_SERVER_ERROR)
            return

        self.send_json({"ok": True, **result})

    def read_request_body(self, max_size: int) -> bytes:
        body_size = int(self.headers.get("Content-Length", "0"))
        if body_size > max_size:
            raise ValidationError("提交内容过大")
        return self.rfile.read(body_size)

    def handle_login(self) -> None:
        ip = self.client_ip()
        user_agent = self.user_agent()
        try:
            body = self.read_request_body(16 * 1024)
            content_type = self.headers.get("Content-Type", "")
            if "application/json" in content_type:
                payload = json.loads(body.decode("utf-8") or "{}")
                password = str(payload.get("password", ""))
            else:
                form = parse_qs(body.decode("utf-8"))
                password = form.get("password", [""])[0]
        except Exception:
            record_login_event(ip, user_agent, False)
            self.send_json({"ok": False, "message": "登录请求无效"}, status=HTTPStatus.BAD_REQUEST)
            return

        if not hmac.compare_digest(password, AUTH_PASSWORD):
            record_login_event(ip, user_agent, False)
            self.send_json({"ok": False, "message": "密码错误"}, status=HTTPStatus.UNAUTHORIZED)
            return

        session_id, expires_at = create_auth_session(ip, user_agent)
        record_login_event(ip, user_agent, True, session_id)
        body = json.dumps({"ok": True}, ensure_ascii=False).encode("utf-8")
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Set-Cookie", self.auth_cookie_header(session_id, expires_at))
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def handle_logout(self) -> None:
        self.expire_current_session()
        body = json.dumps({"ok": True}, ensure_ascii=False).encode("utf-8")
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Set-Cookie", self.clear_auth_cookie_header())
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def record_id_from_path(self, path: str) -> int | None:
        match = re.fullmatch(r"/(?:record|download)/(\d+)", path)
        if not match:
            return None
        return int(match.group(1))

    def send_docx_file(self, path: Path, filename: str) -> None:
        if not path.exists():
            self.send_error(HTTPStatus.NOT_FOUND, "Generated file not found")
            return
        body = path.read_bytes()
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(filename)}")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def handle_import_roster(self) -> None:
        try:
            body_size = int(self.headers.get("Content-Length", "0"))
            if body_size > 8 * 1024 * 1024:
                raise ValidationError("Excel 文件不能超过 8MB")
            payload = json.loads(self.rfile.read(body_size).decode("utf-8"))
            file_name = clean_optional_text(payload.get("fileName"), 180)
            file_data = str(payload.get("fileData", ""))
            if "," in file_data:
                file_data = file_data.split(",", 1)[1]
            if not file_data:
                raise ValidationError("请先选择 Excel 文件")
            result = parse_workbook_entries(base64.b64decode(file_data), file_name)
        except ValidationError as exc:
            self.send_json({"ok": False, "message": str(exc)}, status=HTTPStatus.BAD_REQUEST)
            return
        except Exception as exc:
            self.send_json({"ok": False, "message": f"导入失败：{exc}"}, status=HTTPStatus.INTERNAL_SERVER_ERROR)
            return

        self.send_json({"ok": True, **result})

    def send_json(self, payload: dict, status: HTTPStatus = HTTPStatus.OK) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def main() -> None:
    init_db()
    env_port = os.environ.get("PORT")
    host = os.environ.get("HOST", "0.0.0.0" if env_port else "127.0.0.1")
    port = int(env_port) if env_port else first_free_port(host=host)
    display_host = "127.0.0.1" if host == "0.0.0.0" else host
    url = f"http://{display_host}:{port}"
    server = ThreadingHTTPServer((host, port), AutoFormHandler)
    print(f"思政学分记录表自动生成器已启动：{url}", flush=True)
    print("按 Ctrl+C 停止服务。", flush=True)
    if os.environ.get("OPEN_BROWSER", "1" if not env_port else "0") == "1":
        try:
            webbrowser.open(url)
        except Exception:
            pass
    server.serve_forever()


if __name__ == "__main__":
    main()
